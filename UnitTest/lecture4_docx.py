from docx import Document
from docx.shared import Inches
import os

def readDocument(filePath):
    doc = Document(filePath)
    return doc, doc.core_properties

def writeNewDocument(filePath):
    doc = Document()

    doc.add_heading('Test document form docx', 0)
    doc.save(filePath)

def writeDocument(filePath):
    doc = Document(filePath)

    p = doc.add_paragraph('A plain paragraph having some')
    p.add_run(' bold words').bold = True
    p.add_run(' and italics').italics = True

    doc.add_heading("Let's talk about Python language", 1)
    doc.add_paragraph('First lets see the Python logo', style='ListBullet')

    dir, _ = os.path.split(filePath)
    doc.add_picture(f'{dir}/python.jpg', width=Inches(1.25))

    data = { 'id': 1, 'items': 'apple', 'price': 50 }
    table = doc.add_table(rows=1, cols=3)
    table.style = 'TableGrid'

    headings = table.rows[0].cells
    for i in range(len(table.columns)):
        headings[i].text = list(data.keys())[i]

    row = table.add_row().cells
    for i in range(len(table.columns)):
        row[i].text = str(list(data.values())[i])

    doc.save(filePath)

def generate_document(dir, employee_data, agenda):
    document = Document()

    for emp in employee_data:
        if emp['isDue']:
            name = emp['name']
            department = emp['department']
            
            document.add_heading('Your New Hire Orientation\n', level=1)
            document.add_paragraph('Dear %s' % name)
            document.add_paragraph('Welcome to Google Inc. You have been selected for our new orientation.')
            document.add_paragraph('Based on your department you will go through below sessions:')

            for session in agenda[department]:
                document.add_paragraph(session, style='ListBullet')

            document.add_paragraph('Thanks,\nHR Manager')
            document.save('%s/orientation_%s.docx' % (dir, emp['id']))
